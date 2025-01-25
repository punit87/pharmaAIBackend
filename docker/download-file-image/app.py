import json
from io import BytesIO
import pandas as pd
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

def create_response_with_file(buffer, file_name, mime_type):
    """Creates a response with a file attachment."""
    return {
        "statusCode": 200,
        "headers": {
            "Content-Type": mime_type,
            "Content-Disposition": f'attachment; filename="{file_name}"',
            "Access-Control-Allow-Origin": "*",
        },
        "body": buffer.getvalue().decode('latin-1'),  # Text must be encoded/decoded properly
        "isBase64Encoded": True
    }

def lambda_handler(event, context):
    """Handles file download requests."""
    try:
        data = json.loads(event.get("body", "{}"))
        file_type = data.get("filetype")
        sections = data.get("sections")

        if not sections or not file_type:
            return {
                "statusCode": 400,
                "body": json.dumps({"error": "File type and sections are required!"}),
            }

        buffer = BytesIO()

        if file_type == "word":
            doc = Document()
            for section in sections:
                doc.add_heading(section["subheading"], level=1)
                for row in section.get("table_content", []):
                    doc.add_paragraph(row)
            doc.save(buffer)
            file_name = "output.docx"
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        elif file_type == "pdf":
            c = canvas.Canvas(buffer, pagesize=letter)
            y = 750
            for section in sections:
                c.drawString(50, y, section["subheading"])
                y -= 20
                for row in section.get("table_content", []):
                    row_content = " ".join(row) if isinstance(row, list) else row
                    c.drawString(50, y, row_content)
                    y -= 15
                    if y < 50:
                        c.showPage()
                        y = 750  # Reset y position after a page break
            c.save()
            file_name = "output.pdf"
            mime_type = "application/pdf"

        elif file_type == "text":
            text_content = ""
            for section in sections:
                text_content += f"{section['subheading']}\n"
                for row in section.get("table_content", []):
                    text_content += "\t".join(row) + "\n"
            buffer.write(text_content.encode())
            buffer.seek(0)
            file_name = "output.txt"
            mime_type = "text/plain"

        elif file_type == "excel":
            excel_data = []
            for section in sections:
                section_name = section["subheading"]
                for index, row in enumerate(section.get("table_content", [])):
                    section_column = section_name if index == 0 else ""
                    excel_data.append([section_column, "\t".join(row)])
            df = pd.DataFrame(excel_data, columns=["Section Name", "Table Content"])
            df.to_excel(buffer, index=False)
            file_name = "output.xlsx"
            mime_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

        else:
            return {
                "statusCode": 400,
                "body": json.dumps({"error": "Unsupported file type"}),
            }

        buffer.seek(0)
        return create_response_with_file(buffer, file_name, mime_type)

    except Exception as e:
        return {
            "statusCode": 500,
            "body": json.dumps({"error": f"Error processing file: {str(e)}"}),
        }
